"""
YAMA AI — Document Management API Routes
Upload, analyze, and manage documents.
"""

import os
import uuid
import hashlib
from typing import List, Optional
from datetime import datetime
from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from sqlalchemy.orm import Session
from pydantic import BaseModel

from app.db.database import get_db
from app.db.models import Document, Case, DocumentType
from app.core.config import settings

router = APIRouter(prefix="/documents", tags=["Documents"])

# Upload directory
UPLOAD_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(__file__)))), "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ============ SCHEMAS ============

class DocumentResponse(BaseModel):
    id: int
    doc_uid: str
    case_id: Optional[int]
    filename: str
    original_filename: str
    file_size: Optional[int]
    mime_type: Optional[str]
    document_type: str
    title: Optional[str]
    description: Optional[str]
    extracted_text: Optional[str]
    ocr_processed: bool
    ai_analysis: Optional[str]
    detected_entities: Optional[str]
    relevant_laws: Optional[str]
    uploaded_at: datetime
    analyzed_at: Optional[datetime]

    class Config:
        from_attributes = True


class DocumentUpdate(BaseModel):
    title: Optional[str] = None
    description: Optional[str] = None
    document_type: Optional[str] = None


# ============ HELPER FUNCTIONS ============

def get_file_hash(file_content: bytes) -> str:
    """Generate SHA-256 hash of file content."""
    return hashlib.sha256(file_content).hexdigest()


async def extract_text_from_file(file_path: str, mime_type: str) -> str:
    """Extract text from uploaded file."""
    text = ""
    
    try:
        if mime_type == "text/plain":
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        
        elif mime_type == "application/pdf":
            try:
                import PyPDF2
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() or ""
            except ImportError:
                text = "[PDF extraction requires PyPDF2 library]"
        
        elif mime_type in ["application/msword", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"]:
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                text = "[DOCX extraction requires python-docx library]"
    
    except Exception as e:
        text = f"[Error extracting text: {str(e)}]"
    
    return text


# ============ ROUTES ============

@router.post("", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
async def upload_document(
    file: UploadFile = File(...),
    case_uid: Optional[str] = Form(None),
    document_type: str = Form("other"),
    title: Optional[str] = Form(None),
    description: Optional[str] = Form(None),
    db: Session = Depends(get_db)
):
    """Upload a document, optionally linking to a case."""
    
    # Validate file type
    allowed_types = [
        "application/pdf",
        "image/jpeg", "image/png", "image/gif", "image/webp",
        "application/msword",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
    ]
    
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400, 
            detail=f"File type not allowed. Allowed: PDF, DOC, DOCX, TXT, JPEG, PNG, GIF"
        )
    
    # Validate file size (max 50MB)
    content = await file.read()
    if len(content) > 50 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File size exceeds 50MB limit")
    
    # Get case if provided
    case_id = None
    if case_uid:
        case = db.query(Case).filter(Case.case_uid == case_uid).first()
        if not case:
            raise HTTPException(status_code=404, detail="Case not found")
        case_id = case.id
    
    # Generate unique filename
    doc_uid = str(uuid.uuid4())
    file_ext = os.path.splitext(file.filename)[1]
    filename = f"{doc_uid}{file_ext}"
    file_path = os.path.join(UPLOAD_DIR, filename)
    
    # Save file
    with open(file_path, "wb") as f:
        f.write(content)
    
    # Extract text
    extracted_text = await extract_text_from_file(file_path, file.content_type)
    
    # Create document record
    document = Document(
        doc_uid=doc_uid,
        case_id=case_id,
        filename=filename,
        original_filename=file.filename,
        file_path=file_path,
        file_size=len(content),
        mime_type=file.content_type,
        document_type=DocumentType(document_type) if document_type in [e.value for e in DocumentType] else DocumentType.OTHER,
        title=title or file.filename,
        description=description,
        extracted_text=extracted_text[:50000] if extracted_text else None,  # Limit text storage
        ocr_processed=False,
    )
    
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return {
        "id": document.id,
        "doc_uid": document.doc_uid,
        "case_id": document.case_id,
        "filename": document.filename,
        "original_filename": document.original_filename,
        "file_size": document.file_size,
        "mime_type": document.mime_type,
        "document_type": document.document_type.value if document.document_type else "other",
        "title": document.title,
        "description": document.description,
        "extracted_text": document.extracted_text[:500] if document.extracted_text else None,
        "ocr_processed": document.ocr_processed,
        "ai_analysis": document.ai_analysis,
        "detected_entities": document.detected_entities,
        "relevant_laws": document.relevant_laws,
        "uploaded_at": document.uploaded_at,
        "analyzed_at": document.analyzed_at,
    }


@router.get("", response_model=List[DocumentResponse])
async def list_documents(
    case_uid: Optional[str] = None,
    document_type: Optional[str] = None,
    limit: int = 50,
    offset: int = 0,
    db: Session = Depends(get_db)
):
    """List documents with optional filters."""
    query = db.query(Document)
    
    if case_uid:
        case = db.query(Case).filter(Case.case_uid == case_uid).first()
        if case:
            query = query.filter(Document.case_id == case.id)
    
    if document_type:
        query = query.filter(Document.document_type == DocumentType(document_type))
    
    documents = query.order_by(Document.uploaded_at.desc()).offset(offset).limit(limit).all()
    
    result = []
    for doc in documents:
        result.append({
            "id": doc.id,
            "doc_uid": doc.doc_uid,
            "case_id": doc.case_id,
            "filename": doc.filename,
            "original_filename": doc.original_filename,
            "file_size": doc.file_size,
            "mime_type": doc.mime_type,
            "document_type": doc.document_type.value if doc.document_type else "other",
            "title": doc.title,
            "description": doc.description,
            "extracted_text": doc.extracted_text[:500] if doc.extracted_text else None,
            "ocr_processed": doc.ocr_processed,
            "ai_analysis": doc.ai_analysis,
            "detected_entities": doc.detected_entities,
            "relevant_laws": doc.relevant_laws,
            "uploaded_at": doc.uploaded_at,
            "analyzed_at": doc.analyzed_at,
        })
    
    return result


@router.get("/{doc_uid}", response_model=DocumentResponse)
async def get_document(doc_uid: str, db: Session = Depends(get_db)):
    """Get a specific document."""
    doc = db.query(Document).filter(Document.doc_uid == doc_uid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return {
        "id": doc.id,
        "doc_uid": doc.doc_uid,
        "case_id": doc.case_id,
        "filename": doc.filename,
        "original_filename": doc.original_filename,
        "file_size": doc.file_size,
        "mime_type": doc.mime_type,
        "document_type": doc.document_type.value if doc.document_type else "other",
        "title": doc.title,
        "description": doc.description,
        "extracted_text": doc.extracted_text,
        "ocr_processed": doc.ocr_processed,
        "ai_analysis": doc.ai_analysis,
        "detected_entities": doc.detected_entities,
        "relevant_laws": doc.relevant_laws,
        "uploaded_at": doc.uploaded_at,
        "analyzed_at": doc.analyzed_at,
    }


@router.post("/{doc_uid}/analyze")
async def analyze_document(doc_uid: str, db: Session = Depends(get_db)):
    """Trigger AI analysis of a document."""
    doc = db.query(Document).filter(Document.doc_uid == doc_uid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    if not doc.extracted_text:
        raise HTTPException(status_code=400, detail="No text available for analysis")
    
    # TODO: Integrate with AI analysis service
    # For now, return a placeholder
    from app.services.legal_llm import LegalLLMService
    
    try:
        llm_service = LegalLLMService()
        analysis_prompt = f"""Analyze this legal document and provide:
1. Document Summary
2. Key Legal Issues Identified
3. Relevant Indian Laws/Sections
4. Recommended Actions

Document Content:
{doc.extracted_text[:8000]}
"""
        analysis = await llm_service.analyze_situation(analysis_prompt)
        
        doc.ai_analysis = analysis.get("analysis", "Analysis not available")
        doc.relevant_laws = str(analysis.get("relevant_laws", []))
        doc.analyzed_at = datetime.utcnow()
        
        db.commit()
        
        return {
            "status": "success",
            "doc_uid": doc_uid,
            "analysis": doc.ai_analysis,
            "relevant_laws": doc.relevant_laws,
        }
    
    except Exception as e:
        return {
            "status": "error",
            "doc_uid": doc_uid,
            "error": str(e),
        }


@router.delete("/{doc_uid}", status_code=status.HTTP_204_NO_CONTENT)
async def delete_document(doc_uid: str, db: Session = Depends(get_db)):
    """Delete a document."""
    doc = db.query(Document).filter(Document.doc_uid == doc_uid).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file from disk
    if os.path.exists(doc.file_path):
        os.remove(doc.file_path)
    
    db.delete(doc)
    db.commit()
    return None
